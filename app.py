import streamlit as st
import pdfplumber
from docx import Document
import io
from rank_bm25 import BM25Okapi
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai
from docx import Document as DocWriter
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import os


GEMINI_KEY = os.getenv("GEMINI_KEY")
DEEPSEEK_KEY = os.getenv("DEEPSEEK_KEY")
GROQ_KEY = os.getenv("GROQ_KEY")
OPENROUTER_KEY = os.getenv("OPENROUTER_KEY")
HF_TOKEN = os.getenv("HF_TOKEN")


genai.configure(api_key=GEMINI_KEY)


# ---------- BASIC HELPERS ----------
def chunk_text(pages, chunk_size=500, overlap=100):
    chunks = []
    current = ""
    for page_num, page in enumerate(pages):
        current += (page or "") + " "
        if len(current) >= chunk_size:
            chunks.append({"text": current.strip(), "page": page_num + 1})
            current = current[-overlap:]
    if current.strip():
        chunks.append({"text": current.strip(), "page": len(pages)})
    return chunks


def count_words(text):
    try:
        return len(text.split())
    except:
        return 0


st.set_page_config(page_title="PDF Q&A Tool", layout="wide")
st.title("📚 PDF Question Answer Tool")
st.write("Upload a book PDF and a questions file. The tool will find answers and write them to a new document.")


# ---------- 1: Upload ----------
st.header("1️⃣ Upload files")

pdf_file = st.file_uploader("Upload your PDF book", type=["pdf"])
question_file = st.file_uploader("Upload your questions (DOCX)", type=["docx"])

pdf_text_pages = []
questions = []


# ---------- PDF ----------
if pdf_file is not None:
    st.write("📂 PDF uploaded:", pdf_file.name, "-", pdf_file.size, "bytes")
    progress = st.progress(0)
    try:
        pdf_bytes = pdf_file.getvalue()
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            total = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pdf_text_pages.append(text)
                progress.progress(int(((i + 1) / total) * 100))
        progress.empty()
    except Exception as e:
        st.error(f"❌ PDF read error: {e}")


# ---------- DOCX ----------
if question_file is not None:
    st.write("📂 DOCX uploaded:", question_file.name, "-", question_file.size, "bytes")
    try:
        doc_bytes = question_file.getvalue()
        doc = Document(io.BytesIO(doc_bytes))
        raw_questions = []

        for p in doc.paragraphs:
            if p.text.strip():
                raw_questions.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt and txt not in raw_questions:
                        raw_questions.append(txt)

        questions = []
        keywords = ("Explain", "Discuss", "Describe", "Define", "Evaluate", "What", "How")

        for q in raw_questions:
            q_clean = q.strip()
            if len(q_clean) < 25:
                continue
            if q_clean.startswith(keywords):
                questions.append(q_clean)

        questions = questions[:6]

    except Exception as e:
        st.error(f"❌ DOCX read error: {e}")


# ---------- COUNTS ----------
st.write(f"📄 Pages loaded from PDF: **{len(pdf_text_pages)}**")
st.write(f"❓ Questions detected: **{len(questions)}**")


# ---------- 2: Build search index ----------
bm25 = None
tfidf = None
tfidf_matrix = None
chunks = []

if len(pdf_text_pages) > 0:
    st.header("2️⃣ Build search index")
    progress = st.progress(0)
    chunks = chunk_text(pdf_text_pages)
    texts = [c["text"] for c in chunks]

    bm25 = BM25Okapi([t.split() for t in texts])
    tfidf = TfidfVectorizer().fit(texts)
    tfidf_matrix = tfidf.transform(texts)

    progress.progress(100)
    progress.empty()


def search_answer(question, top_k=4):
    bm25_scores = bm25.get_scores(question.split())
    top_idx = sorted(range(len(bm25_scores)),
                     key=lambda i: bm25_scores[i],
                     reverse=True)[:top_k]

    q_vec = tfidf.transform([question])
    cos_scores = cosine_similarity(q_vec, tfidf_matrix[top_idx]).flatten()
    ranked_idx = [top_idx[i] for i in cos_scores.argsort()[::-1]]
    return [chunks[i] for i in ranked_idx]


# =================================================
#      MODEL CALLERS
# =================================================
def gen_gemini(prompt):
    try:
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)
        return response.text
    except:
        return None



def gen_deepseek(prompt):
    try:
        r = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers={"Authorization": f"Bearer {DEEPSEEK_KEY}"},
            json={
                "model": "deepseek-chat",
                "messages": [{"role": "user", "content": prompt}]
            }
        )
        return r.json()["choices"][0]["message"]["content"]
    except:
        return None


def gen_groq(prompt):
    try:
        r = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_KEY}"},
            json={
                "model": "mixtral-8x7b-32768",
                "messages": [{"role": "user", "content": prompt}]
            }
        )
        return r.json()["choices"][0]["message"]["content"]
    except:
        return None


def gen_openrouter(prompt):
    try:
        r = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {OPENROUTER_KEY}"},
            json={
                "model": "meta-llama/llama-3-8b-instruct",
                "messages": [{"role": "user", "content": prompt}]
            }
        )
        return r.json()["choices"][0]["message"]["content"]
    except:
        return None


# ---------- HUMAN REWRITE ----------
def hf_paraphrase(text):
    try:
        r = requests.post(
            "https://api-inference.huggingface.co/models/tuner007/pegasus_paraphrase",
            headers={"Authorization": f"Bearer {HF_TOKEN}"},
            json={"inputs": text}
        )
        return r.json()[0]["generated_text"]
    except:
        return text


def student_rewrite(text):
    prompt = f"""
Rewrite naturally like a capable university student.

- meaning must remain the same
- formal but not robotic
- avoid repetition
- keep the length similar

TEXT:
{text}
"""
    try:
        result = gen_deepseek(prompt)
        return result or text
    except:
        return text


# =================================================
#   🔥 EXPANDER (ALREADY EXISTS — KEEP)
# =================================================
def expand_answer(existing, question, context, word_limit):
    current = existing.strip()
    loops = 0

    while count_words(current) < word_limit and loops < 4:
        prompt = f"""
You will EXPAND an answer WITHOUT changing it.

Rules:
- ONLY append new text after it
- No rewriting or deleting
- Use context only
- Formal student tone
- Add enough to reach {word_limit} words

Original:
{current}

Context:
{context}

Write ONLY new text to append:
"""
        piece = gen_deepseek(prompt) or gen_groq(prompt) or gen_openrouter(prompt)
        if not piece:
            break

        current = current + "\n\n" + piece.strip()
        loops += 1

    return current


# ==========================================================
#   ⭐ MAX LIMIT CONTROL (UNCHANGED)
# ==========================================================
def shrink_answer(existing, min_words, max_words):
    target = int((min_words + max_words) / 2)

    text = existing
    loops = 0

    while count_words(text) > max_words and loops < 4:
        prompt = f"""
Summarize to about {target} words.

Rules:
- keep meaning and key points
- do NOT add anything new
- single flowing essay
- formal student tone

Text:
{text}
"""
        result = gen_groq(prompt) or gen_deepseek(prompt) or gen_openrouter(prompt)
        if not result:
            break

        text = result.strip()
        loops += 1

    return text


# =================================================
#      FINAL GENERATOR (UNCHANGED)
# =================================================
def generate_answer(question, context, word_limit):

    base_prompt = f"""
Write ONE connected answer.

Question:
{question}

Use ONLY the context:

{context}

Rules:
- combined essay, no headings
- do not restate question
- minimum {word_limit} words
- explain clearly
"""
    MODELS = [
        ("Gemini", gen_gemini),
        ("DeepSeek", gen_deepseek),
        ("Groq", gen_groq),
        ("OpenRouter", gen_openrouter),
    ]

    draft = None
    st.write("➡ **Content Generation Phase started…**")

    for round_ in range(2):
        for label, fn in MODELS:
            st.write(f"🔎 Generating with **{label}** …")
            try:
                draft = fn(base_prompt)
                if draft and count_words(draft) > int(word_limit * 0.6):
                    break
            except:
                continue
        if draft:
            break

    if not draft:
        return "Answer generation failed due to multiple API limits. Please try again later."

    st.write("✍️ **Human rewrite phase…**")
    draft = hf_paraphrase(draft)
    draft = student_rewrite(draft)

    if count_words(draft) < word_limit:
        draft = expand_answer(draft, question, context, word_limit)

    return draft.strip()


# ---------- DOCX HELPERS ----------
def set_style(run, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold


#  UPDATED — FORMAT LIKE PICTURE-1 
def save_answers_to_docx(qa_pairs, output_name="answers.docx"):
    doc = DocWriter()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # TITLE OUTSIDE TABLE
    title = doc.add_paragraph("ASSIGNMENT")
    title.runs[0].bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # DETAILS TABLE
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("NAME", name),
        ("PROGRAM", program),
        ("ROLL NO.", roll),
        ("SESSION", session_value),
        ("SEMESTER", semester),
        ("COURSE CODE", course_code),
        ("COURSE NAME", course_name),
    ]

    for label, value in rows:
        r = table.add_row().cells
        r[0].text = label
        r[1].text = value or ""

        for run in r[0].paragraphs[0].runs:
            set_style(run, bold=True)

        for run in r[1].paragraphs[0].runs:
            set_style(run)

    doc.add_paragraph()

    midpoint = len(qa_pairs) // 2

    def write_block(i, q, a):
        p = doc.add_paragraph()
        r1 = p.add_run(f"QUESTION {i}: ")
        set_style(r1, bold=True)
        r2 = p.add_run(q)
        set_style(r2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p2 = doc.add_paragraph()
        r3 = p2.add_run("ANSWER:")
        set_style(r3, bold=True)

        answer_text = a.replace("\n- ", "\n• ")
        p3 = doc.add_paragraph(answer_text)
        for r in p3.runs:
            set_style(r)
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    s1 = doc.add_paragraph("SET 1")
    set_style(s1.runs[0], bold=True)
    s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for i, (q, a) in enumerate(qa_pairs[:midpoint], start=1):
        write_block(i, q, a)

    s2 = doc.add_paragraph("SET 2")
    set_style(s2.runs[0], bold=True)
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for i, (q, a) in enumerate(qa_pairs[midpoint:], start=midpoint + 1):
        write_block(i, q, a)

    doc.save(output_name)
    return output_name


# ---------- 3: Word limits ----------
st.header("3️⃣ Word limit")
word_limit = st.number_input(
    "Enter desired MINIMUM words per answer",
    min_value=50,
    max_value=2000,
    value=200,
    step=10
)

max_words = st.number_input(
    "Enter MAXIMUM words per answer",
    min_value=80,
    max_value=6000,
    value=400,
    step=10
)

# ---------- 4: Assignment Details (NEW UI FIELDS) ----------
st.header("4️⃣ Assignment Details")

name = st.text_input("Student Name")
program = st.text_input("Program")
roll = st.text_input("Roll Number")
session_value = st.text_input("Session")
semester = st.text_input("Semester")
course_code = st.text_input("Course Code")
course_name = st.text_input("Course Name")


# ---------- 5: Run tool ----------
st.header("5️⃣ Run the tool")

if st.button("Process and Generate Answers"):

    if not pdf_file or not question_file:
        st.warning("Please upload everything first.")
    else:
        st.success("Generating answers...")

        max_total_words = 12 * 480
        per_answer_cap = max_total_words // max(len(questions), 1)
        effective_limit = min(word_limit, per_answer_cap)

        qa_pairs = []
        progress = st.progress(0)

        for i, q in enumerate(questions):
            try:
                ctx = search_answer(q, top_k=6)
                context = "\n\n".join([c["text"] for c in ctx])
            except:
                context = ""

            answer = generate_answer(q, context, effective_limit)

            if count_words(answer) > max_words:
                st.write("✂️ Trimming oversized answer…")
                answer = shrink_answer(answer, effective_limit, max_words)

            qa_pairs.append((q, answer))
            progress.progress(int(((i + 1) / len(questions)) * 100))

        file_path = save_answers_to_docx(qa_pairs)

        st.success("Done — formatted DOCX ready.")
        st.download_button(
            "📥 Download answers",
            data=open(file_path, "rb").read(),
            file_name="answers.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


